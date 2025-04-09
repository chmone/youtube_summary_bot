import os
import json
import logging
from datetime import datetime, timezone
import requests
import xml.etree.ElementTree as ET
import re
import urllib.parse
import time
import argparse
from youtube_transcript_api import YouTubeTranscriptApi
import yt_dlp
from docx import Document
import dropbox
import spacy
from apscheduler.schedulers.blocking import BlockingScheduler
from dotenv import load_dotenv

# --- Configuration Loading ---
load_dotenv() # Load environment variables from .env file for sensitive keys

CONFIG_FILE = 'config.json'
OPENROUTER_API_BASE = "https://openrouter.ai/api/v1"

def load_config():
    """Loads configuration from config.json and environment variables."""
    try:
        with open(CONFIG_FILE, 'r') as f:
            config = json.load(f)

        # Load sensitive keys from environment variables, overriding if necessary
        config['YOUTUBE_API_KEY'] = os.getenv('YOUTUBE_API_KEY', config.get('YOUTUBE_API_KEY'))
        config['OPENROUTER_API_KEY'] = os.getenv('OPENROUTER_API_KEY')
        config['DROPBOX_ACCESS_TOKEN'] = os.getenv('DROPBOX_ACCESS_TOKEN', config.get('DROPBOX_ACCESS_TOKEN'))

        # Validate required keys
        if not config.get('OPENROUTER_API_KEY'):
             logging.warning("OPENROUTER_API_KEY not found in environment variables. AI features will be disabled.")
        # Note: Dropbox token is checked later during initialization/upload

        return config
    except FileNotFoundError:
        logging.error(f"Configuration file {CONFIG_FILE} not found.")
        exit(1)
    except json.JSONDecodeError:
        logging.error(f"Error decoding JSON from {CONFIG_FILE}.")
        exit(1)

config = load_config()

# --- Logging Setup ---
log_file = config.get('LOG_FILE', 'logs/app.log')
os.makedirs(os.path.dirname(log_file), exist_ok=True)

# Define handlers with UTF-8 encoding specified
file_handler = logging.FileHandler(log_file, encoding='utf-8')
stream_handler = logging.StreamHandler()
# Attempt to set encoding for stream handler (console) - might depend on terminal capabilities
try:
    stream_handler.setStream(open(stream_handler.stream.fileno(), mode='w', encoding='utf-8', buffering=1))
except Exception:
    # Fallback if setting encoding on existing stream fails (e.g., in some environments)
    pass 

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s',
                    # Pass the configured handlers
                    handlers=[
                        file_handler,
                        stream_handler
                    ])

# --- API Client Initialization & Key Check ---
openrouter_api_key = config.get('OPENROUTER_API_KEY')
dropbox_access_token = config.get('DROPBOX_ACCESS_TOKEN')
dbx = None
nlp = None

try:
    if dropbox_access_token:
        dbx = dropbox.Dropbox(dropbox_access_token)
    else:
        logging.warning("Dropbox access token not found in config or environment variables. Upload disabled.")

    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        logging.error("spaCy 'en_core_web_sm' model not found. Please run: python -m spacy download en_core_web_sm")
        exit(1)

except Exception as e:
    logging.error(f"Error during API client initialization: {e}")
    exit(1)

# --- OpenRouter Helper Function ---
def call_openrouter_api(messages, model=None):
    """Calls the OpenRouter Chat Completions API."""
    if not openrouter_api_key:
        logging.error("OpenRouter API key is missing. Cannot make API call.")
        return None

    if not model:
        model = config.get('OPENROUTER_MODEL', 'google/gemini-2.5-pro-exp-03-25')  # Updated default # Default if not specified

    headers = {
        "Authorization": f"Bearer {openrouter_api_key}",
        # Optional headers for identification:
        # "HTTP-Referer": "YOUR_SITE_URL",
        # "X-Title": "YouTubeSummaryBot",
        "Content-Type": "application/json"
    }
    payload = {
        "model": model,
        "messages": messages,
        # Add other params like temperature if needed
        # "temperature": 0.7,
    }
    url = f"{OPENROUTER_API_BASE}/chat/completions"

    try:
        response = requests.post(url, headers=headers, json=payload)
        response.raise_for_status() # Raise HTTPError for bad responses (4xx or 5xx)
        response_data = response.json()

        # Extract content
        if response_data.get("choices") and len(response_data["choices"]) > 0:
            message_data = response_data["choices"][0].get("message", {})
            content = message_data.get("content")
            if content:
                logging.debug(f"OpenRouter response received from {model}")
                return content.strip()
            else:
                logging.error(f"OpenRouter response missing content in choices[0].message. Full response: {response_data}")
                return None
        else:
            # Log the full response data when choices are missing or empty
            logging.error(f"OpenRouter response missing choices field or empty choices. Full response: {response_data}")
            return None

    except requests.exceptions.RequestException as e:
        logging.error(f"Error calling OpenRouter API: {e}")
        # Log detailed error response if available
        try:
            if response:
                error_details = response.json()
                logging.error(f"OpenRouter error details: {error_details}")
            else:
                logging.error("No response object available to extract error details.")
        except Exception as json_e: # Handle cases where response is not available or not JSON
            logging.error(f"Could not parse error response as JSON: {json_e}")
        return None
    except Exception as e:
        logging.exception(f"An unexpected error occurred during OpenRouter API call: {e}") # Log stack trace
        return None

# --- New Transcript Analysis Function ---
def parse_analysis_response(response_text):
    """Parses the structured markdown response from the LLM."""
    sections = {
        'main_points': 'Not found in response.',
        'tools': 'Not found in response.',
        'principles': 'Not found in response.',
        'vocab': 'Not found in response.'
    }
    if not response_text:
        return sections

    # Define patterns for each section header (case-insensitive, allows optional whitespace)
    patterns = {
        'main_points': r'\*\*Main Useful Points:\*\*(.*?)'
                       r'(?=\*\*Emerging Tools or Ideas Mentioned:|\*\*Workflow or Coding Principles:|\*\*Technical Vocabulary and Concepts:|\Z)',
        'tools': r'\*\*Emerging Tools or Ideas Mentioned:\*\*(.*?)'
                 r'(?=\*\*Main Useful Points:|\*\*Workflow or Coding Principles:|\*\*Technical Vocabulary and Concepts:|\Z)',
        'principles': r'\*\*Workflow or Coding Principles:\*\*(.*?)'
                      r'(?=\*\*Main Useful Points:|\*\*Emerging Tools or Ideas Mentioned:|\*\*Technical Vocabulary and Concepts:|\Z)',
        'vocab': r'\*\*Technical Vocabulary and Concepts:\*\*(.*?)'
                 r'(?=\*\*Main Useful Points:|\*\*Emerging Tools or Ideas Mentioned:|\*\*Workflow or Coding Principles:|\Z)'
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, response_text, re.IGNORECASE | re.DOTALL)
        if match:
            content = match.group(1).strip()
            if content:
                sections[key] = content
            else:
                sections[key] = "(Section found but empty)" # Indicate if section header exists but content is missing
            logging.debug(f"Parsed section '{key}':\n{sections[key][:100]}...") # Log start of parsed content
        else:
            logging.warning(f"Could not find section '{key}' in LLM response.")

    return sections

def analyze_transcript(transcript):
    """Analyzes the transcript using a single OpenRouter call with structured prompting."""
    if not openrouter_api_key:
        logging.warning("OpenRouter API key not configured. Cannot analyze transcript.")
        # Return a dictionary with error messages
        return {
            'main_points': 'Analysis unavailable (API key missing).',
            'tools': 'Analysis unavailable (API key missing).',
            'principles': 'Analysis unavailable (API key missing).',
            'vocab': 'Analysis unavailable (API key missing).'
        }

    logging.info("Analyzing transcript via OpenRouter with structured prompt...")

    system_prompt = (
        "You are an AI assistant trained to extract high-value insights from YouTube video transcripts. "
        "Your job is to provide an educational, concise, and structured summary that helps the user deeply understand key concepts, emerging trends, and useful ideas from the video. "
        "The video is posted by a technical content creator focused on AI, programming, startups, and productivity."
    )
    user_prompt = (
        "Here is a transcript from a recent video:\n\n" # Use \n for line breaks
        f"{transcript}\n\n" # Use \n for line breaks
        "Please provide a structured output in the following format:\n\n" # Use \n for line breaks
        "**Main Useful Points:**\n"
        "- Bullet points of key insights, learnings, strategies, or opinions. Be concise but thoughtful.\n"
        "- Include context where helpful. Avoid generic takeaways.\n\n"
        "**Emerging Tools or Ideas Mentioned:**\n"
        "- List tools, frameworks, or models referenced (with a 1-line description).\n\n"
        "**Workflow or Coding Principles:**\n"
        "- Extract any workflow tips, coding best practices, or structured development techniques.\n\n"
        "**Technical Vocabulary and Concepts:**\n"
        "- Format: Word: \"Definition\"\n"
        "- Include a brief example sentence or use-case from the transcript if possible."
    )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]

    analysis_response_text = call_openrouter_api(messages)

    if analysis_response_text:
        logging.info("Analysis received from OpenRouter. Parsing response...")
        parsed_analysis = parse_analysis_response(analysis_response_text)
        return parsed_analysis
    else:
        logging.error("Failed to get analysis response from OpenRouter.")
        # Return a dictionary indicating failure
        return {
            'main_points': 'Analysis failed (API Error).',
            'tools': 'Analysis failed (API Error).',
            'principles': 'Analysis failed (API Error).',
            'vocab': 'Analysis failed (API Error).'
        }

# --- Text Cleaning Helper ---
def clean_llm_output(text):
    """Performs basic cleaning on text received from LLM, including removing markdown bold."""
    if not isinstance(text, str):
        return ""
    
    # Remove markdown bold markers
    cleaned = text.replace('**', '')
    
    # Normalize newlines and strip overall whitespace
    cleaned = cleaned.replace('\r\n', '\n').replace('\r', '\n').strip()
    
    # Split into lines, strip each line, and remove empty lines
    lines = cleaned.split('\n')
    cleaned_lines = [line.strip() for line in lines if line.strip()]
    
    # Rejoin the cleaned lines
    return '\n'.join(cleaned_lines)

# --- Core Functions ---

def get_recent_video_infos_from_rss(channel_id, max_count=5): # Fetch a few by default
    """
    Fetches metadata for recent videos from the channel's RSS feed.
    Assumes feed lists newest videos first. 
    Returns a list of dictionaries [{'id': ..., 'title': ..., 'published': ...}], 
    up to max_count, in the order found in the feed (newest first).
    Returns an empty list on error.
    """
    if not channel_id:
        logging.error("YouTube Channel ID is missing.")
        return []
    feed_url = f"https://www.youtube.com/feeds/videos.xml?channel_id={channel_id}"
    logging.info(f"Checking RSS feed for recent videos: {feed_url}")
    namespaces = {'atom': 'http://www.w3.org/2005/Atom', 'yt': 'http://www.youtube.com/xml/schemas/2015'}
    videos = []

    try:
        response = requests.get(feed_url, timeout=15)
        response.raise_for_status()
        root = ET.fromstring(response.content)

        entries = root.findall('atom:entry', namespaces)
        logging.info(f"Found {len(entries)} entries in RSS feed.")

        # Process entries in the order they appear (should be newest first)
        for i, entry in enumerate(entries):
            if len(videos) >= max_count:
                break # Stop processing if we hit the desired count

            video_id_tag = entry.find('yt:videoId', namespaces)
            title_tag = entry.find('atom:title', namespaces)
            published_tag = entry.find('atom:published', namespaces)

            if video_id_tag is not None and title_tag is not None and published_tag is not None \
               and video_id_tag.text and title_tag.text and published_tag.text:
                
                video_id = video_id_tag.text
                video_title = title_tag.text
                published_str = published_tag.text
                published_dt = None
                try:
                    published_dt = datetime.fromisoformat(published_str).astimezone(timezone.utc)
                except ValueError:
                    logging.warning(f"Could not parse date '{published_str}' for video {video_id}.")
                
                video_data = {
                    'id': video_id,
                    'title': video_title,
                    'published': published_dt
                }
                videos.append(video_data)
                
                # Log details of the very first video found
                if i == 0:
                    logging.info(f"First video entry found in feed: ID={video_id}, Title='{video_title}', Published={published_str}")
            else:
                 logging.warning(f"Skipping RSS entry #{i+1} due to missing id, title, or published date.")

        # Return the list as found (newest first)
        logging.info(f"Returning {len(videos)} video infos (newest first).")
        return videos

    except requests.exceptions.Timeout:
        logging.error(f"Timeout error fetching RSS feed: {feed_url}")
        return []
    except requests.exceptions.RequestException as e:
        logging.error(f"Error fetching RSS feed: {e}")
        return []
    except ET.ParseError as e:
        logging.error(f"Error parsing XML feed: {e}")
        try:
            logging.debug(f"XML content (start): {response.text[:500]}...")
        except NameError: 
             pass 
        return []
    except Exception as e:
        logging.exception(f"An unexpected error occurred fetching/parsing RSS feed: {e}")
        return []

def get_last_processed_video_id():
    """Reads the last processed video ID from a file with added logging."""
    filepath = config['LAST_PROCESSED_VIDEO_ID_FILE']
    logging.info(f"Attempting to read last processed ID from: {filepath}")
    try:
        with open(filepath, 'r') as f:
            video_id = f.read().strip()
            logging.info(f"Successfully read last processed ID: {video_id}")
            return video_id
    except FileNotFoundError:
        logging.warning(f"State file not found: {filepath}. Will process the latest video as new.")
        return None
    except Exception as e:
        logging.error(f"Error reading state file {filepath}: {e}")
        return None # Treat other errors as if no ID was found

def save_last_processed_video_id(video_id):
    """Saves the last processed video ID to a file with added logging."""
    filepath = config['LAST_PROCESSED_VIDEO_ID_FILE']
    logging.info(f"Attempting to save last processed ID '{video_id}' to: {filepath}")
    try:
        # Ensure the directory exists within the volume
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        with open(filepath, 'w') as f:
            f.write(video_id)
        logging.info(f"Successfully saved last processed ID: {video_id}")
    except Exception as e:
        logging.error(f"Error writing state file {filepath}: {e}")

def download_video(video_url, video_id):
    """Downloads the video using yt-dlp (optional based on config)."""
    if not config.get('DOWNLOAD_VIDEOS', False):
        logging.info("Video download is disabled in config.")
        return None

    output_path = f"downloads/{video_id}.%(ext)s"
    ydl_opts = {
        'format': 'best', # Or specify a format
        'outtmpl': output_path,
        'quiet': True,
    }
    try:
        logging.info(f"Attempting to download video: {video_url}")
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info_dict = ydl.extract_info(video_url, download=True)
            downloaded_path = ydl.prepare_filename(info_dict)
            logging.info(f"Video downloaded to: {downloaded_path}")
            return downloaded_path
    except Exception as e:
        logging.error(f"Error downloading video {video_url}: {e}")
        return None

def get_transcript(video_id):
    """Fetches the transcript using youtube-transcript-api."""
    try:
        logging.info(f"Fetching transcript for video ID: {video_id}")
        transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
        transcript_text = " ".join([item['text'] for item in transcript_list])
        logging.info(f"Transcript fetched successfully for video ID: {video_id}")
        return transcript_text
    except Exception as e:
        logging.error(f"Could not retrieve transcript for video ID {video_id}: {e}")
        return None

def generate_document(video_id, video_title, video_url, analysis, transcript, published_date=None):
    """Generates the output document (.docx) using the video's publish date in the filename."""
    output_format = config.get('OUTPUT_FORMAT', 'docx').lower()
    
    # Determine the date string for the filename
    if published_date and isinstance(published_date, datetime):
        date_str = published_date.strftime("%Y-%m-%d")
        logging.debug(f"Using publish date for filename: {date_str}")
    else:
        logging.warning("Publish date not available or invalid, using current date for filename.")
        date_str = datetime.now().strftime("%Y-%m-%d")

    safe_title = re.sub(r'[\\/*?"<>|:]', '', video_title)
    safe_title = re.sub(r'\s+', ' ', safe_title).strip()
    safe_title = safe_title[:60]
    if not safe_title: safe_title = "Untitled_Video"
    
    filename_base = f"{date_str}_{safe_title}_Summary" # Use determined date_str
    output_dir = "outputs/"
    os.makedirs(output_dir, exist_ok=True)

    filepath = None
    logging.info(f"Generating {output_format} document with chat-style formatting...")

    if output_format == 'docx':
        filepath = os.path.join(output_dir, f"{filename_base}.docx")
        doc = Document()
        
        # Add main title (still using a heading for the overall document title)
        doc.add_heading(f"Analysis for: {video_title}", level=1)
        doc.add_paragraph(f"Video URL: {video_url}\n") # Add a blank line after URL

        # --- Section Helper Function ---
        def add_section(title_text, content_text):
            # Add title paragraph and make it bold
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(title_text)
            title_run.bold = True
            # Add content paragraph (cleaned)
            content_paragraph = doc.add_paragraph(clean_llm_output(content_text))
            # Add a blank paragraph for spacing after the content
            doc.add_paragraph() 

        # --- Add Sections using Helper ---
        add_section('Main Useful Points', analysis.get('main_points', 'Not provided.'))
        add_section('Emerging Tools or Ideas Mentioned', analysis.get('tools', 'Not provided.'))
        add_section('Workflow or Coding Principles', analysis.get('principles', 'Not provided.'))
        add_section('Technical Vocabulary and Concepts', analysis.get('vocab', 'Not provided.'))
        
        # --- Add Transcript Section ---
        # Use bold paragraph for transcript title
        transcript_title_paragraph = doc.add_paragraph()
        transcript_title_run = transcript_title_paragraph.add_run("Full Transcript")
        transcript_title_run.bold = True
        # Add cleaned transcript content
        doc.add_paragraph(clean_llm_output(transcript))

        try:
            doc.save(filepath)
            logging.info(f"Document saved to: {filepath}")
        except Exception as e:
            logging.error(f"Error saving document {filepath}: {e}")
            filepath = None

    else:
        logging.error(f"Unsupported output format: {output_format}")
        return None

    return filepath

def upload_to_dropbox(filepath):
    """Uploads the specified file to Dropbox."""
    if not dbx:
        logging.error("Dropbox client not initialized. Cannot upload.")
        return
    if not filepath or not os.path.exists(filepath):
        logging.error(f"Filepath invalid or file does not exist: {filepath}")
        return

    filename = os.path.basename(filepath)
    dropbox_path = os.path.join(config.get('DROPBOX_UPLOAD_PATH', '/DailySummaries/'), filename).replace('\\', '/') # Ensure forward slashes

    try:
        logging.info(f"Uploading {filename} to Dropbox path: {dropbox_path}")
        with open(filepath, 'rb') as f:
            dbx.files_upload(f.read(), dropbox_path, mode=dropbox.files.WriteMode('overwrite'))
        logging.info(f"Successfully uploaded {filename} to Dropbox.")
    except dropbox.exceptions.ApiError as e:
        logging.error(f"Dropbox API error during upload: {e}")
    except Exception as e:
        logging.error(f"Error uploading file {filename} to Dropbox: {e}")

# --- Main Processing Logic ---

def process_video(video_info):
    """Processes a single video given its info dictionary.
       Returns True if successful, False otherwise.
    """
    video_id = video_info['id']
    video_title = video_info['title']
    published_date = video_info.get('published') # Get the publish date
    video_url = f"https://www.youtube.com/watch?v={video_id}"
    logging.info(f"--- Processing video: '{video_title}' ({video_id}) --- Published: {published_date}")

    # 1. Download Video (Optional)
    download_video(video_url, video_id)

    # 2. Get Transcript
    transcript = get_transcript(video_id)
    if not transcript:
        logging.error(f"Failed to get transcript for {video_id}. Skipping video.")
        return False

    # 3. Analyze Transcript
    analysis_data = analyze_transcript(transcript)
    # Basic check if analysis failed
    if analysis_data.get('main_points', '').startswith('Analysis failed'):
        logging.error(f"Analysis failed for {video_id}. Skipping document generation.")
        return False # Indicate failure

    # 4. Generate Document - Pass the published_date
    output_filepath = generate_document(video_id, video_title, video_url, analysis_data, transcript, published_date)

    # 5. Upload to Cloud
    if output_filepath and dbx:
        upload_to_dropbox(output_filepath)
    elif output_filepath:
         logging.warning("Dropbox upload skipped (no token/client init failed).")
    else:
        logging.error("Doc generation failed, cannot upload.")
        # Consider if doc generation failure should stop processing or just be logged
        return False # Treat doc failure as overall failure for this video

    logging.info(f"--- Finished processing video: '{video_title}' ({video_id}) ---")
    return True

def process_latest_video_if_new():
    """Checks the latest video from RSS (assuming newest first) and processes if new."""
    logging.info("Checking for the single latest video...")
    channel_id = config.get('YOUTUBE_CHANNEL_ID')
    if not channel_id:
        logging.error("YOUTUBE_CHANNEL_ID not set.")
        return

    # Fetch recent videos (newest first)
    recent_videos = get_recent_video_infos_from_rss(channel_id, max_count=5) # Get up to 5
    if not recent_videos:
        logging.warning("Could not determine any video info from RSS feed.")
        return

    # The first video in the list should be the actual latest one
    latest_video_info = recent_videos[0] 
    latest_video_id = latest_video_info['id']
    video_title = latest_video_info['title']
    
    last_processed_id = get_last_processed_video_id()
    logging.info(f"Comparing Latest ID from feed '{latest_video_id}' ('{video_title}') with Last Processed ID '{last_processed_id}'")

    if latest_video_id == last_processed_id:
        logging.info(f"No new video found. IDs match.")
        return

    logging.info(f"New video detected! ID '{latest_video_id}' != Last Processed ID '{last_processed_id}'")
    
    success = process_video(latest_video_info)

    if success:
        logging.info(f"Video processing successful. Preparing to save ID: {latest_video_id}")
        save_last_processed_video_id(latest_video_id)
    else:
         logging.warning(f"Processing failed for new video {latest_video_id}, not updating last processed ID.")

# --- New Backfill Function ---
def backfill_last_n_videos(count):
    """Fetches and processes the specified number of most recent videos, regardless of processed state."""
    logging.warning(f"--- Starting Backfill Mode for last {count} videos --- ")
    logging.warning("NOTE: Backfill mode does NOT update the last_processed_id file.")
    
    channel_id = config.get('YOUTUBE_CHANNEL_ID')
    if not channel_id:
        logging.error("YOUTUBE_CHANNEL_ID not set. Cannot run backfill.")
        return

    # Fetch slightly more than requested in case some fail?
    # Fetching max (e.g., 15) and slicing is safer.
    recent_videos = get_recent_video_infos_from_rss(channel_id, max_count=15)
    if not recent_videos:
        logging.error("Could not retrieve any videos from RSS feed for backfill.")
        return

    # Get the actual last 'count' videos from the list (newest first for processing)
    videos_to_process = recent_videos[-count:] # Slice the newest 'count' videos
    videos_to_process.reverse() # Process newest first? Or oldest first? Let's do newest first.
    
    logging.info(f"Attempting to backfill {len(videos_to_process)} videos (up to {count} requested)...")

    processed_count = 0
    failed_count = 0
    for video_info in videos_to_process:
        try:
            success = process_video(video_info)
            if success:
                processed_count += 1
            else:
                failed_count += 1
            # Optional: Add a short delay between videos in backfill mode?
            # time.sleep(5) # e.g., 5 seconds
        except Exception as e:
            failed_count += 1
            logging.exception(f"Unhandled exception processing video {video_info.get('id', 'N/A')} during backfill: {e}")
           
    logging.warning(f"--- Finished Backfill Mode --- ")
    logging.info(f"Backfill summary: Successfully processed = {processed_count}, Failed = {failed_count}")

# --- Scheduler ---

def run_scheduler():
    scheduler = BlockingScheduler()
    polling_hours = config.get('POLLING_INTERVAL_HOURS', 6)
    logging.info(f"Scheduler starting. Check RSS every {polling_hours} hrs.")
    try:
        # Initial check uses the function that compares with last processed ID
        process_latest_video_if_new()
    except Exception as e:
        logging.exception("Unhandled exception during initial run.")
    
    # Subsequent checks also only process if new
    scheduler.add_job(process_latest_video_if_new, 'interval', hours=polling_hours)
    logging.info(f"Scheduled job to run every {polling_hours} hours.")
    try:
        scheduler.start()
    except (KeyboardInterrupt, SystemExit):
        logging.info("Scheduler stopped by user.")
    except Exception as e:
        logging.exception("Scheduler failed unexpectedly.")

if __name__ == "__main__":
    logging.info("Application starting...")
    
    # --- Argument Parsing ---
    parser = argparse.ArgumentParser(description="YouTube Video Summarizer Bot.")
    parser.add_argument(
        "--backfill", 
        type=int, 
        metavar="N",
        help="Process the last N videos from the RSS feed, ignoring the last processed ID."
    )
    args = parser.parse_args()

    # --- Main Execution Logic ---
    if args.backfill:
        if args.backfill > 0:
            backfill_last_n_videos(args.backfill)
        else:
            logging.error("Value for --backfill must be a positive integer.")
    else:
        # Run the scheduler (normal mode)
        run_scheduler() 